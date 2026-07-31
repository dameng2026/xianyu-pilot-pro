"""Generate test Word/Excel files for KB upload mode testing."""
import os
from docx import Document
from openpyxl import Workbook

OUT_DIR = os.path.dirname(os.path.abspath(__file__))


def make_word():
    doc = Document()
    doc.add_heading("客服知识库 - 图书教材类问答手册", level=1)
    doc.add_paragraph("本手册整理了图书教材类目下常见的客服问答，可用于 AI 客服培训与知识库录入参考。")

    doc.add_heading("一、商品真伪与质量", level=2)
    doc.add_paragraph("Q：这本书是正版吗？")
    doc.add_paragraph("A：是的，本店所售图书均为正版，支持专柜验货。我们与出版社直接合作采购，所有书籍均带有出版社防伪标识，假一赔十。")

    doc.add_paragraph("Q：书籍有破损怎么办？")
    doc.add_paragraph("A：收到商品后请在 48 小时内检查，如有破损、缺页、印刷质量问题，请拍照联系客服，我们将在 24 小时内为您免费补寄或退款。")

    doc.add_heading("二、物流发货", level=2)
    doc.add_paragraph("Q：什么时候发货？")
    doc.add_paragraph("A：工作日 15:00 前付款的订单当天发货，15:00 后付款的订单次日发货。节假日发货时间顺延，江浙沪皖一般 1-2 天送达，其他地区 3-5 天。")

    doc.add_paragraph("Q：支持自选快递吗？")
    doc.add_paragraph("A：默认发中通/圆通，如需顺丰请补差价 10 元，新疆西藏青海等偏远地区需补运费 15-25 元，下单后联系客服修改。")

    doc.add_heading("三、退换货政策", level=2)
    doc.add_paragraph("Q：可以七天无理由退货吗？")
    doc.add_paragraph("A：支持 7 天无理由退货，但请保证书籍未使用、未书写、包装完整。塑封书籍拆封后不支持无理由退货，质量问题除外。")

    doc.add_paragraph("Q：退货邮费谁出？")
    doc.add_paragraph("A：质量问题导致的退货，邮费由本店承担；非质量问题退货由买家承担回寄邮费。请勿使用顺丰或到付，否则可能被拒收。")

    doc.add_heading("四、教材版本与适用年级", level=2)
    doc.add_paragraph("Q：这本教材是哪个版本的？")
    doc.add_paragraph("A：商品详情页已标注版本（人教版/北师大版/苏教版等），请仔细核对。如不确定学校使用版本，建议咨询任课老师后再下单。")

    doc.add_paragraph("Q：适合几年级使用？")
    doc.add_paragraph("A：每本教材封面均有年级标识，商品标题也注明了适用年级。若您购买的是同步练习册，请选择与孩子当前学期一致的版本。")

    doc.add_heading("五、批发与优惠", level=2)
    doc.add_paragraph("Q：批量购买有优惠吗？")
    doc.add_paragraph("A：单笔订单满 10 本享 95 折，满 30 本享 9 折，满 100 本请联系客服申请团购价。教师凭教师证购买可叠加 5% 优惠。")

    doc.add_paragraph("Q：有优惠券吗？")
    doc.add_paragraph("A：每月 1 号、15 号会发放店铺优惠券，关注店铺即可领取。老客户回购可联系客服领取专属回归券，满 50 减 5 元。")

    out = os.path.join(OUT_DIR, "客服知识库_图书教材_测试.docx")
    doc.save(out)
    print(f"已生成 Word: {out}")


def make_excel():
    wb = Workbook()
    ws = wb.active
    ws.title = "客服问答"
    headers = ["问题", "回答", "分类", "标签"]
    ws.append(headers)
    rows = [
        ["这本书是正版吗？", "是的，本店所售图书均为正版，支持专柜验货，假一赔十。", "图书教材/正版验证", "正版,验货"],
        ["支持七天无理由退货吗？", "支持 7 天无理由退货，请保证书籍未使用、包装完整。塑封拆封后不支持无理由退货。", "图书教材/退换货", "退货,7天无理由"],
        ["什么时候发货？", "工作日 15:00 前付款当天发货，之后次日发货。江浙沪皖 1-2 天，其他地区 3-5 天。", "图书教材/物流发货", "发货,物流"],
        ["新疆能发货吗？", "新疆、西藏、青海等偏远地区需补运费 15-25 元，下单后联系客服修改运费。", "图书教材/物流发货", "偏远地区,运费"],
        ["这本教材是人教版吗？", "商品详情页已标注版本，请仔细核对。如不确定，建议咨询任课老师后再下单。", "图书教材/版本咨询", "人教版,版本"],
        ["批量购买有优惠吗？", "单笔满 10 本 95 折，满 30 本 9 折，满 100 本联系客服申请团购价。", "图书教材/批发优惠", "批发,团购,优惠"],
        ["书籍破损怎么办？", "48 小时内拍照联系客服，24 小时内免费补寄或退款。", "图书教材/售后", "破损,补寄,售后"],
        ["退货邮费谁出？", "质量问题本店承担，非质量问题买家承担。请勿顺丰或到付。", "图书教材/退换货", "退货,邮费"],
        ["适合几年级？", "封面有年级标识，标题也注明适用年级。同步练习册请选择与孩子学期一致的版本。", "图书教材/版本咨询", "年级,练习册"],
        ["有优惠券吗？", "每月 1 号、15 号发放店铺券，关注店铺领取。老客户回购可领回归券满 50 减 5。", "图书教材/优惠活动", "优惠券,回归券"],
    ]
    for r in rows:
        ws.append(r)

    # 设置列宽
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 60
    ws.column_dimensions['C'].width = 22
    ws.column_dimensions['D'].width = 18

    out = os.path.join(OUT_DIR, "客服知识库_图书教材_测试.xlsx")
    wb.save(out)
    print(f"已生成 Excel: {out}")


def make_markdown():
    content = """# 客服知识库 - 数码电子类问答

## 一、商品质量

**Q：手机是全新未拆封的吗？**
A：是的，本店所有手机均为全新国行正品，未拆封带官方保修。支持苹果官网查询序列号。

**Q：屏幕有坏点怎么办？**
A：收到商品 7 天内如发现坏点、亮点超过 3 个，可申请换货。请提供屏幕检测视频，我们承担来回运费。

## 二、物流与发货

**Q：发什么快递？多久到？**
A：默认发顺丰，工作日 17:00 前付款当天发货。一线城市 1-2 天，二三线城市 2-4 天。

**Q：可以自提吗？**
A：暂不支持自提，所有订单均通过快递发货，可签收后联系客服开具电子发票。

## 三、售后保修

**Q：保修多久？**
A：国行手机享受苹果官方 1 年保修，本店额外赠送 6 个月店铺保修，覆盖屏幕、电池、主板。

**Q：进水了保修吗？**
A：进水、摔坏、私拆等人为损坏不在保修范围内，可联系客服付费维修，更换屏幕约 400-800 元。
"""
    out = os.path.join(OUT_DIR, "客服知识库_数码电子_测试.md")
    with open(out, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"已生成 Markdown: {out}")


if __name__ == "__main__":
    make_word()
    make_excel()
    make_markdown()
    print("全部测试文件生成完成")
